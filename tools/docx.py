"""Word (.docx) 工具：读取、新建、追加、替换文本。

依赖 python-docx（间接依赖 lxml，a-Shell 首次安装需现场编译，约几分钟）。
import 是惰性的：未安装时工具会返回友好提示而不是崩溃。
"""

from pathlib import Path
from tools import Tool


def _resolve(gate, p: str) -> Path:
    return gate.sandbox_root.joinpath(p).resolve()


def _docx():
    try:
        import docx
        return docx
    except ImportError:
        return None


def make_docx_tools(gate) -> list[Tool]:
    def docx_read(args):
        p = _resolve(gate, args["path"])
        if not gate.read(p):
            return "拒绝：路径越界或密钥文件"
        d = _docx()
        if d is None:
            return "未安装 python-docx：请先 pip install python-docx"
        try:
            doc = d.Document(str(p))
        except Exception as e:  # noqa: BLE001
            return f"读取失败：{e}"
        lines = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(lines) or "（空文档）"

    def docx_create(args):
        p = _resolve(gate, args["path"])
        if p.exists():
            if not gate.overwrite(p):
                return "拒绝：未确认覆盖"
        elif not gate.write(p):
            return "拒绝：未确认写入"
        d = _docx()
        if d is None:
            return "未安装 python-docx"
        doc = d.Document()
        for line in str(args["content"]).split("\n"):
            s = line.rstrip()
            if not s.strip():
                continue
            if s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            else:
                doc.add_paragraph(s)
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        return f"已创建 {args['path']}"

    def docx_append(args):
        p = _resolve(gate, args["path"])
        if not p.exists():
            return f"文件不存在：{args['path']}"
        if not gate.overwrite(p):
            return "拒绝：未确认覆盖"
        d = _docx()
        if d is None:
            return "未安装 python-docx"
        doc = d.Document(str(p))
        doc.add_paragraph(str(args["text"]))
        doc.save(str(p))
        return f"已追加到 {args['path']}"

    def docx_replace(args):
        p = _resolve(gate, args["path"])
        if not p.exists():
            return f"文件不存在：{args['path']}"
        if not gate.overwrite(p):
            return "拒绝：未确认覆盖"
        d = _docx()
        if d is None:
            return "未安装 python-docx"
        doc = d.Document(str(p))
        old, new = args["old"], args["new"]
        count = 0
        for para in doc.paragraphs:
            if old not in para.text:
                continue
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
        doc.save(str(p))
        return f"已替换 {count} 处" if count else "未找到要替换的文本"

    return [
        Tool("docx_read", "读取 Word(.docx) 文档的文本内容（含表格）",
             {"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}, docx_read),
        Tool("docx_create", "新建 Word 文档。content 传多行字符串，'# '开头为一级标题、'## '为二级标题，其余为段落",
             {"type": "object", "properties": {"path": {"type": "string"}, "content": {"type": "string"}},
              "required": ["path", "content"]}, docx_create),
        Tool("docx_append", "向 Word 文档末尾追加一段文字",
             {"type": "object", "properties": {"path": {"type": "string"}, "text": {"type": "string"}},
              "required": ["path", "text"]}, docx_append),
        Tool("docx_replace", "替换 Word 文档中的文本（段落 run 级别）",
             {"type": "object", "properties": {"path": {"type": "string"}, "old": {"type": "string"}, "new": {"type": "string"}},
              "required": ["path", "old", "new"]}, docx_replace),
    ]
